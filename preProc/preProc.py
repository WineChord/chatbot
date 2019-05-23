# -*- coding: utf-8 -*-
from docx import Document
from processDoc import *

def preProc():
    """
    main function
    :para:
    :return:

    labels:
    identifier: e.g '2.3.1'
    index: e.g '2.3.1 ALM-0x12000001 进风口温度过高(Chassis,轻微告警)'
    explanation:  告警解释
    impact_on_sys: 对系统的影响
    causes: 可能原因
    solutions: 处理步骤

    put each of these labels into documents respectively
    """
    faulty_part_records, sensor_records = getData()
    mode1 = './故障部件模式/故障部件模式'
    mode2 = './传感器模式/传感器模式'
    outputDoc(faulty_part_records, mode1)
    outputDoc(sensor_records, mode2)

def getData():
    """
    get data from secDoc
    :para
    :return (faulty_part_records, sensor_records)
    """
    name1 = './srcDoc/故障部件模式告警.docx'
    name2 = './srcDoc/传感器模式告警.docx'
    return proc(name1), proc(name2)

def outputDoc(r, mode):
    """
    output results of records to documents
    :para r (records)
    :return None
    """
    # generate faulty_part document
    d_identifier = Document()
    d_index = Document()
    d_explanation = Document()
    d_impact = Document()
    d_causes = Document()
    d_solutions = Document()
    for i in range(len(r)):
        d_identifier.add_paragraph(r[i].identifier)
        d_index.add_paragraph(r[i])
        d_explanation.add_paragraph(r[i].explanation)
        d_impact.add_paragraph(r[i].impact_on_sys)
        d_causes.add_paragraph(r[i].causes)
        d_solutions.add_paragraph(r[i].solutions)
    d_identifier.save(mode+'_告警序号.docx')
    d_index.save(mode+'_告警内容.docx')
    d_explanation.save(mode+'_告警解释.docx')
    d_impact.save(mode+'_对系统的影响.docx')
    d_causes.save(mode+'_可能原因.docx')
    d_solutions.save(mode+'_处理步骤.docx')

def deleteHT(argv):
    """
    delete headers and trailers
    """
#    argv = '故障部件模式告警.docx'
    document = Document(argv)
    # lines in document
    l = [paragraph.text for paragraph in document.paragraphs]

    newl = []
    pattern = '文档版本 24 (2018-08-08)'
    i = 0
    while i < len(l):
        if l[i] == pattern:  # delete headers and trailers
            i += 6
        newl.append(l[i])
        i += 1

#    newd = Document()
#    for i in range(len(newl)):
#        newd.add_paragraph(newl[i])
#    newd.save('testoutput.docx')  # save middle result
    return newl

def proc(argv):
    """
    essential part for document parse
    :para document name
    :return records
    """
    partl = deleteHT(argv) # for test
    records = parseDoc(partl)
    return records

def uut():
    name2 = './srcDoc/传感器模式告警.docx'
    return deleteHT(name2)


if __name__ == '__main__':
    preProc()
